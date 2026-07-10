"""Document discovery and extraction helpers."""

from __future__ import annotations

import hashlib
import re
from email import policy
from email.parser import BytesParser
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import BinaryIO

from docx import Document
from openpyxl import load_workbook

from backend.config import settings
from backend.schemas import DocumentAsset, DocumentCategory, DocumentPacket


CATEGORY_FOLDERS: dict[DocumentCategory, str] = {
    DocumentCategory.service_agreement: "sample_contract_style_service_agreements",
    DocumentCategory.invoice: "invoices",
    DocumentCategory.billing_evidence: "raw_billing_evidence",
    DocumentCategory.payment_record: "raw_payment_records",
}


def normalize_customer_hint(file_name: str) -> str | None:
    stem = Path(file_name).stem.lower()
    tokens = re.split(r"[_\-\s]+", stem)
    noise = {
        "sample",
        "contract",
        "inv",
        "invoice",
        "style",
        "service",
        "agreement",
        "raw",
        "payment",
        "record",
        "pending",
        "receipt",
        "advice",
        "application",
        "statement",
        "july",
        "export",
        "ledger",
        "thread",
        "note",
        "notes",
        "snapshot",
        "submittal",
        "based",
        "time",
        "materials",
        "over",
        "cap",
        "managed",
        "services",
        "outcome",
        "milestone",
        "arc",
        "rrc",
        "tm",
        "ms",
        "mile",
        "out",
        "001",
        "002",
        "01",
        "02",
        "a",
        "b",
    }
    kept = [token for token in tokens if token and token not in noise and not token.isdigit()]
    return "_".join(kept[:3]) or None


def load_packet_from_corpus(billing_type: str, documents_root: Path | None = None) -> DocumentPacket:
    root = documents_root or settings.documents_root
    packet = DocumentPacket(billing_type=billing_type)
    for category, folder_name in CATEGORY_FOLDERS.items():
        folder = root / folder_name / billing_type
        assets = _load_assets_from_folder(folder, category, billing_type)
        _assign_assets(packet, category, assets)
    return packet


def load_packet_from_uploads(
    billing_type: str,
    service_agreements: list[Path],
    invoices: list[Path],
    billing_evidence: list[Path],
    payment_records: list[Path],
) -> DocumentPacket:
    packet = DocumentPacket(billing_type=billing_type)
    packet.service_agreements = [_asset_from_path(path, DocumentCategory.service_agreement, billing_type) for path in service_agreements]
    packet.invoices = [_asset_from_path(path, DocumentCategory.invoice, billing_type) for path in invoices]
    packet.billing_evidence = [_asset_from_path(path, DocumentCategory.billing_evidence, billing_type) for path in billing_evidence]
    packet.payment_records = [_asset_from_path(path, DocumentCategory.payment_record, billing_type) for path in payment_records]
    return packet


def save_uploaded_files(files: list[BinaryIO], target_dir: Path) -> list[Path]:
    target_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    for file_obj in files:
        name = Path(getattr(file_obj, "name", "uploaded_file")).name
        path = target_dir / name
        data = file_obj.getvalue() if hasattr(file_obj, "getvalue") else file_obj.read()
        path.write_bytes(data)
        paths.append(path)
    return paths


def uploaded_packet_context() -> TemporaryDirectory[str]:
    return TemporaryDirectory(prefix="billing-agent-uploads-")


def packet_to_prompt(packet: DocumentPacket) -> str:
    sections = [
        f"Billing type: {packet.billing_type}",
        _category_prompt("SERVICE AGREEMENTS", packet.service_agreements),
        _category_prompt("INVOICES", packet.invoices),
        _category_prompt("RAW BILLING EVIDENCE", packet.billing_evidence),
        _category_prompt("PAYMENT RECORDS", packet.payment_records),
    ]
    return "\n\n".join(sections)


def _load_assets_from_folder(folder: Path, category: DocumentCategory, billing_type: str) -> list[DocumentAsset]:
    if not folder.exists():
        return []
    return [
        _asset_from_path(path, category, billing_type)
        for path in sorted(folder.iterdir())
        if path.is_file() and not path.name.startswith(".")
    ]


def _assign_assets(packet: DocumentPacket, category: DocumentCategory, assets: list[DocumentAsset]) -> None:
    if category == DocumentCategory.service_agreement:
        packet.service_agreements = assets
    elif category == DocumentCategory.invoice:
        packet.invoices = assets
    elif category == DocumentCategory.billing_evidence:
        packet.billing_evidence = assets
    elif category == DocumentCategory.payment_record:
        packet.payment_records = assets


def _asset_from_path(path: Path, category: DocumentCategory, billing_type: str) -> DocumentAsset:
    try:
        text = extract_text(path)
        parse_status = "parsed"
    except Exception as exc:  # pragma: no cover - defensive extraction fallback
        text = ""
        parse_status = f"parse_error: {exc}"
    document_id = hashlib.sha1(str(path).encode("utf-8")).hexdigest()[:12]
    return DocumentAsset(
        document_id=document_id,
        category=category,
        file_name=path.name,
        path=str(path),
        extension=path.suffix.lower(),
        billing_type=billing_type,
        customer_hint=normalize_customer_hint(path.name),
        extracted_text=text[: settings.max_document_chars],
        parse_status=parse_status,
    )


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix == ".eml":
        return _extract_eml(path)
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(errors="replace")
    return f"Unsupported file type {suffix}. File path: {path}"


def _extract_docx(path: Path) -> str:
    doc = Document(path)
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for table_index, table in enumerate(doc.tables, start=1):
        chunks.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _extract_xlsx(path: Path) -> str:
    workbook = load_workbook(path, data_only=True, read_only=True)
    chunks: list[str] = []
    for sheet in workbook.worksheets:
        chunks.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks)


def _extract_eml(path: Path) -> str:
    with path.open("rb") as handle:
        message = BytesParser(policy=policy.default).parse(handle)
    parts = [
        f"Subject: {message.get('subject', '')}",
        f"From: {message.get('from', '')}",
        f"To: {message.get('to', '')}",
        f"Date: {message.get('date', '')}",
    ]
    if message.is_multipart():
        for part in message.walk():
            if part.get_content_type() == "text/plain":
                parts.append(part.get_content())
    else:
        parts.append(message.get_content())
    return "\n".join(str(part) for part in parts if part)


def _category_prompt(title: str, assets: list[DocumentAsset]) -> str:
    if not assets:
        return f"{title}: none provided"
    rendered = [f"{title}:"]
    for asset in assets:
        rendered.append(
            "\n".join(
                [
                    f"- File: {asset.file_name}",
                    f"  Category: {asset.category.value}",
                    f"  Customer hint: {asset.customer_hint or 'unknown'}",
                    f"  Parse status: {asset.parse_status}",
                    "  Extracted text:",
                    _indent(asset.extracted_text or "[no extracted text]", "    "),
                ]
            )
        )
    return "\n".join(rendered)


def _indent(text: str, prefix: str) -> str:
    return "\n".join(prefix + line for line in text.splitlines())
